"""
运维记录簿（YWJLB）统一处理模块
支持三种包类型：01包-金财工程、01包-其他项目、02包项目
"""

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
import os
import logging
from enum import Enum

# 获取模块日志记录器（日志配置由应用入口统一管理）
logger = logging.getLogger(__name__)


class PackageType(Enum):
    """包类型定义"""
    GJGC = "01包-金财工程"  # 01 Package - Jincai Engineering
    QTXM = "01包-其他项目"  # 01 Package - Other Projects
    PKG02 = "02包项目"      # 02 Package Projects


# 定义各包类型对应的字段
PACKAGE_FIELDS = {
    PackageType.GJGC: [
        ("项目名称", "项目名称"),
        ("需求名称", "需求名称"),
        ("使用单位", "使用单位"),
        ("监理单位", "监理单位"),
        ("运维单位", "运维单位"),
        ("提出问题单位", "提出问题单位"),
        ("日期", "日期"),
        ("运维人员", "运维人员"),
        ("系统模块", "系统模块"),
        ("问题类型", "问题类型"),
        ("问题描述", "问题描述"),
        ("处理方法", "处理方法"),
        ("处理结果", "处理结果"),
    ],
    PackageType.QTXM: [
        ("项目名称", "项目名称"),
        ("需求名称", "需求名称"),
        ("使用单位", "使用单位"),
        ("监理单位", "监理单位"),
        ("运维单位", "运维单位"),
        ("日期", "日期"),
        ("运维人员", "运维人员"),
        ("问题类型", "问题类型"),
        ("问题描述", "问题描述"),
        ("处理方法", "处理方法"),
        ("处理结果", "处理结果"),
    ],
    PackageType.PKG02: [
        ("项目名称", "项目名称"),
        ("需求名称", "需求名称"),
        ("使用单位", "使用单位"),
        ("监理单位", "监理单位"),
        ("运维单位", "运维单位"),
        ("日期", "日期"),
        ("运维人员", "运维人员"),
        ("问题类型", "问题类型"),
        ("问题描述", "问题描述"),
        ("处理方法", "处理方法"),
        ("处理结果", "处理结果"),
    ]
}


def set_east_asian_font_for_run(run, font_name):
    """为run对象设置中文字体"""
    try:
        r = run._r
        rPr = r.get_or_add_rPr()
        font = OxmlElement('w:rFonts')
        font.set(qn('w:eastAsia'), font_name)
        rPr.append(font)
    except Exception as e:
        logging.error(f"设置Run对象字体时出错: {str(e)}")


def set_east_asian_font_for_style(style, font_name):
    """为样式对象设置中文字体"""
    try:
        rPr = style._element.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            style._element.insert(0, rPr)
        
        font = OxmlElement('w:rFonts')
        font.set(qn('w:eastAsia'), font_name)
        rPr.append(font)
    except Exception as e:
        logging.error(f"设置Style对象字体时出错: {str(e)}")


def remove_paragraph_border(paragraph):
    """移除段落的边框（包括分隔线）"""
    try:
        pPr = paragraph._p.get_or_add_pPr()
        
        # 检查是否存在边框设置并移除
        if pPr.find(qn('w:pBdr')) is not None:
            pPr.remove(pPr.find(qn('w:pBdr')))
            
        # 明确设置无边界
        pBdr = OxmlElement('w:pBdr')
        
        # 为所有边设置无边界
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            tag = OxmlElement(f'w:{side}')
            tag.set(qn('w:val'), 'none')
            pBdr.append(tag)
            
        pPr.append(pBdr)
    except Exception as e:
        logging.error(f"移除段落边框时出错: {str(e)}")


def set_table_column_width(table, col_index, width):
    """强制设置表格列宽"""
    try:
        # 设置列宽
        table.columns[col_index].width = width
        
        # 遍历该列的所有单元格，确保宽度设置生效
        for row in table.rows:
            cell = row.cells[col_index]
            # 设置单元格宽度
            cell.width = width
            
            # 通过XML直接设置宽度（更可靠的方法）
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width.twips)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            
    except Exception as e:
        logging.error(f"设置表格列宽时出错: {str(e)}")


def create_signature_table(doc):
    """添加独立的签字区域表格（1行2列）"""
    try:
        # 创建1行2列的表格用于签字区域
        sign_table = doc.add_table(rows=1, cols=2, style='Table Grid')
        sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 定义列宽
        LABEL_COL_WIDTH = Cm(3.8)
        CONTENT_COL_WIDTH = Cm(12.87)  # 16.67 - 3.8 = 12.87 

        # 设置列宽
        set_table_column_width(sign_table, 0, LABEL_COL_WIDTH)
        set_table_column_width(sign_table, 1, CONTENT_COL_WIDTH)
        
        # 设置签字区域内容
        left_cell = sign_table.cell(0, 0)
        left_paragraph = left_cell.paragraphs[0]
        left_paragraph.text = "\n运维人员\n签字"
        left_paragraph.runs[0].font.bold = True
        left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        
        right_cell = sign_table.cell(0, 1)
        right_paragraph = right_cell.paragraphs[0]
        right_paragraph.text = "\n\n\n    年    月    日"
        right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # 设置表格行高，留足签字空间
        sign_table.rows[0].height = Cm(3)
        
        logging.debug("签字区域表格创建成功")
        
    except Exception as e:
        logging.error(f"创建签字区域表格时出错: {str(e)}")


def create_content_table(doc, label, content):
    """创建单个内容表格，对日期进行格式化处理"""
    try:
        table = doc.add_table(rows=1, cols=2, style='Table Grid')
        
        # 定义列宽
        LABEL_COL_WIDTH = Cm(3.8)
        CONTENT_COL_WIDTH = Cm(12.87)  # 16.67 - 3.8 = 12.87
        
        # 设置列宽
        set_table_column_width(table, 0, LABEL_COL_WIDTH)
        set_table_column_width(table, 1, CONTENT_COL_WIDTH)
        
        # 设置内容，对日期进行格式化处理
        label_cell = table.cell(0, 0)
        content_cell = table.cell(0, 1)
        
        label_cell.text = label
        
        # 处理日期格式，只显示年月日
        if label == "日期":
            if pd.notna(content):
                # 尝试将内容转换为日期并格式化
                if isinstance(content, pd.Timestamp):
                    content_str = content.strftime("%Y年%m月%d日")
                else:
                    try:
                        # 尝试解析字符串格式的日期
                        date_obj = pd.to_datetime(content)
                        content_str = date_obj.strftime("%Y年%m月%d日")
                    except Exception:
                        content_str = str(content)
            else:
                content_str = ""
        else:
            content_str = str(content) if pd.notna(content) else ""
            
        content_cell.text = content_str
        
        # 设置标签列的对齐方式（上下左右居中）和字体加粗
        label_paragraph = label_cell.paragraphs[0]
        label_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        label_cell.vertical_alignment = 1  # 垂直居中 (WD_CELL_VERTICAL_ALIGNMENT.CENTER = 1)
        
        # 设置标签列字体加粗
        for run in label_paragraph.runs:
            run.font.bold = True
        
        # 根据标签类型设置内容列的对齐方式和行高
        content_paragraph = content_cell.paragraphs[0]
        # 需要居中的字段
        center_fields = [
            "项目名称", "需求名称", "使用单位", "监理单位", "运维单位", 
            "提出问题单位", "日期", "运维人员", "系统模块"
        ]
        
        if label in center_fields:
            # 这些字段内容上下左右居中
            content_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            content_cell.vertical_alignment = 1  # 垂直居中
            # 设置固定行高为0.61厘米
            table.rows[0].height = Cm(0.61)
        else:
            # 其他字段内容左对齐，垂直居中
            content_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            content_cell.vertical_alignment = 1  # 垂直居中
            # 设置默认行高
            table.rows[0].height = Cm(2.6)
        
        logging.debug(f"内容表格创建成功: {label}")
        return table
        
    except Exception as e:
        logging.error(f"创建内容表格时出错 ({label}): {str(e)}")
        raise


def format_date_for_filename(date_value):
    """将日期格式化为文件名可用的格式"""
    try:
        if pd.isna(date_value) or date_value == '':
            return ""
        
        # 如果是pandas Timestamp对象
        if isinstance(date_value, pd.Timestamp):
            return date_value.strftime("%Y%m%d")
        
        # 尝试解析其他格式的日期
        try:
            date_obj = pd.to_datetime(date_value)
            return date_obj.strftime("%Y%m%d")
        except Exception:
            # 如果无法解析，返回原始值的安全版本
            safe_date = "".join([c for c in str(date_value) if c.isalnum()])
            return safe_date[:8] if len(safe_date) > 8 else safe_date
            
    except Exception as e:
        logging.warning(f"格式化日期时出错: {str(e)}")
        return ""


def create_word_document(row, idx, package_type: PackageType):
    """创建单个Word文档并返回文档对象"""
    try:
        doc = Document()
        
        # 设置文档页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
        
        # 添加表头"运维记录表"并移除下划线和分隔线
        heading = doc.add_heading(level=0)
        run = heading.add_run("运维记录表")
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(16)
        run.font.name = '仿宋_GB2312'
        run.font.underline = False  # 明确移除下划线
        set_east_asian_font_for_run(run, '仿宋_GB2312')
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 移除标题样式中的下划线和边框（分隔线）
        if heading.style.font.underline:
            heading.style.font.underline = False
        
        # 移除段落边框（包括可能的分隔线）
        remove_paragraph_border(heading)
        
        # 设置文档字体（全局默认）
        style = doc.styles['Normal']
        font = style.font
        font.name = '仿宋_GB2312'
        font.size = Pt(14)
        font.underline = False  # 确保默认样式也没有下划线
        set_east_asian_font_for_style(style, '仿宋_GB2312')
        
        # 获取该包类型对应的字段列表
        fields_config = PACKAGE_FIELDS[package_type]
        
        # 创建各个内容表格
        for label, key in fields_config:
            content = row.get(label, '')
            create_content_table(doc, label, content)
        
        # 添加签字区域表格
        create_signature_table(doc)
        
        logging.debug(f"第{idx}页文档创建成功 (包类型: {package_type.value})")
        return doc
        
    except Exception as e:
        error_msg = f"创建第{idx}页文档时发生错误: {str(e)}"
        logging.error(error_msg)
        raise Exception(error_msg)


def save_word_document(doc, idx, sheet_name, date_value, output_dir="运维记录文档"):
    """保存Word文档，增加sheet_name和output_dir参数
    
    Args:
        doc: Word文档对象
        idx: 文档索引号
        sheet_name: 工作表名称
        date_value: 日期值
        output_dir: 输出文件夹路径（默认为"运维记录文档"）
    """
    try:
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 处理工作表名称中的非法字符
        safe_sheet_name = "".join([c for c in sheet_name if c not in r'/\:*?"<>|'])

        # 格式化日期为文件名可用格式
        date_str = format_date_for_filename(date_value)
        
        # 文件名格式：运维记录_表名_第x页_日期.docx
        if date_str:
            file_path = os.path.join(output_dir, f"运维记录_{safe_sheet_name}_第{idx}页_{date_str}.docx")
        else:
            # 如果没有日期，则使用原格式
            file_path = os.path.join(output_dir, f"运维记录_{safe_sheet_name}_第{idx}页.docx")

        doc.save(file_path)
        logging.info(f"成功生成: {file_path}")
        return True
        
    except PermissionError:
        error_msg = f"保存第{idx}页文档时权限不足，请检查文件是否被占用或是否有写入权限"
        logging.error(error_msg)
        raise Exception(error_msg)
    except Exception as e:
        error_msg = f"保存第{idx}页文档时发生错误: {str(e)}"
        logging.error(error_msg)
        raise Exception(error_msg)


def process_excel_file(excel_file: str, package_type: PackageType, output_dir: str = "运维记录文档"):
    """处理Excel文件并生成Word文档
    
    Args:
        excel_file: Excel文件路径
        package_type: 包类型（PackageType枚举值）
        output_dir: 输出文件夹路径（默认为"运维记录文档"）
    """
    try:
        logging.info(f"开始转换Excel数据到Word文档... (包类型: {package_type.value})")
        logging.info(f"输出文件夹: {output_dir}")
        
        if not os.path.exists(excel_file):
            raise FileNotFoundError(f"找不到Excel文件: {excel_file}")
        
        try:
            # 读取所有工作表
            all_sheets = pd.read_excel(excel_file, sheet_name=None)
            logging.info(f"成功读取Excel文件，共{len(all_sheets)}个工作表")
        except Exception as e:
            raise Exception(f"读取Excel文件失败: {str(e)}")
        
        total_success = 0
        total_failed = 0
        
        # 遍历每个工作表
        for sheet_name, df in all_sheets.items():
            logging.info(f"开始处理工作表: {sheet_name}")
            
            if df.empty:
                logging.warning(f"工作表[{sheet_name}]中没有数据，跳过处理")
                continue
                
            data_rows = df.to_dict('records')
            logging.info(f"工作表[{sheet_name}]包含{len(data_rows)}行数据")
            
            # 每个工作表内部从1开始计数
            for sheet_idx, row in enumerate(data_rows, start=1):
                try:
                    logging.debug(f"开始处理工作表[{sheet_name}]的第{sheet_idx}行数据")
                    doc = create_word_document(row, sheet_idx, package_type)
                    # 保存时传入工作表名称、索引、日期和输出文件夹
                    save_word_document(doc, sheet_idx, sheet_name, row.get('日期', ''), output_dir)
                    total_success += 1
                    logging.debug(f"工作表[{sheet_name}]的第{sheet_idx}行数据处理成功")
                except Exception as e:
                    logging.error(f"处理工作表[{sheet_name}]的第{sheet_idx}行数据失败: {str(e)}", exc_info=True)
                    total_failed += 1
                    continue
                    
                if sheet_idx % 5 == 0:
                    logging.info(f"工作表[{sheet_name}]已处理{sheet_idx}/{len(data_rows)}行数据")
        
        logging.info(f"转换完成！成功生成{total_success}个文档，失败{total_failed}个")
        print(f"转换完成！成功生成{total_success}个文档，失败{total_failed}个")
        print(f"文档已保存到: {output_dir}")
        print("详细日志请查看：日志文件.log")
        
        return total_success, total_failed
        
    except Exception as e:
        logging.critical(f"程序执行失败: {str(e)}", exc_info=True)
        print(f"错误: {str(e)}")
        print("详细错误信息请查看日志文件")
        raise
